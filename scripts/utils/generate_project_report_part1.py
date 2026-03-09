"""
Generate B.Tech Final Year Project Report - SOC Copilot
Part 1: Helpers, Front Matter, Chapters 1-4
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS_DIR = os.path.join(OUTPUT_DIR, "docs", "diagrams")

def set_margins(section, top=1, bottom=1, left=1.25, right=1):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_centered(doc, text, size=14, bold=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold
    return p

def add_body(doc, text, size=12, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    if indent: p.paragraph_format.first_line_indent = Inches(0.5)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = FONT; r.font.color.rgb = RGBColor(0,0,0)
    return h

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.clear()
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size)
    return p

def add_image(doc, filename, caption="", width=5.5):
    """Add an image from the diagrams directory with optional caption."""
    img_path = os.path.join(DIAGRAMS_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(8)
            r = cap.add_run(caption)
            r.font.name = FONT; r.font.size = Pt(10); r.italic = True
    else:
        add_body(doc, f"[Image not found: {filename}]", italic=True)

def set_cell_shading(cell, color):
    tc = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color); shd.set(qn("w:val"), "clear")
    tc.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(11); r.font.name = FONT
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(c, "D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(11); r.font.name = FONT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph("")
    return t

def add_page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════
def write_title_page(doc):
    for _ in range(4): doc.add_paragraph("")
    add_centered(doc, "A PROJECT REPORT ON", 16, True, 12)
    add_centered(doc, '"SOC Copilot: An AI-Powered Security Operations Center\nAssistant for Threat Detection, Log Analysis,\nand Incident Response"', 18, True, 18)
    add_centered(doc, "Submitted in partial fulfillment of the requirements\nfor the award of the degree of", 13, space_after=12, space_before=18)
    add_centered(doc, "BACHELOR OF TECHNOLOGY", 16, True, 6)
    add_centered(doc, "in", 13, space_after=6)
    add_centered(doc, "COMPUTER SCIENCE AND ENGINEERING", 14, True, 18)
    add_centered(doc, "Submitted by:", 13, space_after=6, space_before=18)
    for i in range(1,5):
        add_centered(doc, f"[Student {i} Name]  —  [Roll Number {i}]", 12, space_after=4)
    add_centered(doc, "Under the Guidance of", 13, space_after=6, space_before=18)
    add_centered(doc, "[Guide Name], [Designation]", 13, True, 12)
    add_centered(doc, "Department of Computer Science and Engineering", 13, space_after=4, space_before=18)
    add_centered(doc, "[College Name]", 14, True, 4)
    add_centered(doc, "[City, State — Pin Code]", 12, space_after=12)
    add_centered(doc, "Academic Year: 2025–2026", 13, True)

# ══════════════════════════════════════════
# CERTIFICATE, DECLARATION, ACKNOWLEDGEMENT
# ══════════════════════════════════════════
def write_certificate(doc):
    add_page_break(doc)
    add_centered(doc, "CERTIFICATE", 18, True, 18, 12)
    add_body(doc, 'This is to certify that the project entitled "SOC Copilot: An AI-Powered Security Operations Center Assistant for Threat Detection, Log Analysis, and Incident Response" is a bonafide work carried out by the following students in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering from [College Name], affiliated to [University Name], during the academic year 2025–2026.')
    add_body(doc, "[Student 1 Name] — [Roll Number 1]")
    add_body(doc, "[Student 2 Name] — [Roll Number 2]")
    add_body(doc, "[Student 3 Name] — [Roll Number 3]")
    add_body(doc, "[Student 4 Name] — [Roll Number 4]")
    add_body(doc, "The results embodied in this project report have not been submitted to any other University or Institution for the award of any degree or diploma.")
    for _ in range(4): doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("[Guide Name]\n[Designation]\nDepartment of CSE\n[College Name]")
    r.font.name = FONT; r.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("[Head of Department]\nDepartment of CSE\n[College Name]")
    r2.font.name = FONT; r2.font.size = Pt(12)

def write_declaration(doc):
    add_page_break(doc)
    add_centered(doc, "DECLARATION", 18, True, 18, 12)
    add_body(doc, 'We hereby declare that the project entitled "SOC Copilot: An AI-Powered Security Operations Center Assistant for Threat Detection, Log Analysis, and Incident Response" submitted to the Department of Computer Science and Engineering, [College Name], in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology, is a record of original work done by us under the supervision and guidance of [Guide Name], [Designation], Department of CSE, [College Name].')
    add_body(doc, "The results embodied in this project report have not been submitted to any other University or Institution for the award of any degree or diploma. We also declare that no part of this work has been copied from any other source. Any material or idea taken from other work has been properly cited and acknowledged.")
    for _ in range(3): doc.add_paragraph("")
    add_body(doc, "Place: [City]")
    add_body(doc, "Date: [Date]")
    for _ in range(2): doc.add_paragraph("")
    for i in range(1,5):
        add_body(doc, f"[Student {i} Name]")

def write_acknowledgement(doc):
    add_page_break(doc)
    add_centered(doc, "ACKNOWLEDGEMENT", 18, True, 18, 12)
    add_body(doc, "We would like to express our sincere gratitude to all who contributed to the successful completion of this project. First and foremost, we thank our project guide, [Guide Name], [Designation], Department of Computer Science and Engineering, [College Name], for the invaluable guidance, constant encouragement, and constructive suggestions throughout the development of this project.")
    add_body(doc, "We extend our heartfelt thanks to [HOD Name], Head of the Department of Computer Science and Engineering, for providing us with the necessary facilities and support. We are grateful to [Principal Name], Principal, [College Name], for the institutional support.")
    add_body(doc, "We also thank all the faculty members and staff of the Department of Computer Science and Engineering for their support and encouragement. We are thankful to our families and friends for their unwavering support and motivation.")
    add_body(doc, "Finally, we acknowledge the contributions of the open-source community, particularly the developers of Python, Scikit-learn, PyQt6, and the CICIDS2017 dataset creators, whose work made this project possible.")

# ══════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════
def write_abstract(doc):
    add_page_break(doc)
    add_centered(doc, "ABSTRACT", 18, True, 18, 12)
    add_body(doc, "The exponential growth of cyber threats in the modern digital landscape has placed unprecedented pressure on Security Operations Centers (SOCs), which serve as the frontline defense against cyberattacks targeting organizational infrastructure. SOC analysts face the daunting challenge of monitoring, triaging, and responding to thousands of security alerts generated daily from heterogeneous log sources, including firewalls, intrusion detection systems, endpoint agents, and cloud workloads. Traditional Security Information and Event Management (SIEM) systems, while instrumental in log aggregation, rely predominantly on static, rule-based detection mechanisms that are inherently reactive and incapable of adapting to the rapidly evolving threat landscape characterized by advanced persistent threats (APTs), zero-day exploits, polymorphic malware, and sophisticated social engineering campaigns.")
    add_body(doc, 'This project presents "SOC Copilot," an AI-powered Security Operations Center assistant designed to automate threat detection, intelligent log analysis, and incident response through a hybrid machine learning ensemble approach. SOC Copilot integrates two complementary machine learning paradigms: Isolation Forest for unsupervised anomaly detection, which identifies behavioral deviations from established network baselines without requiring labeled training data, and Random Forest for supervised multi-class attack classification, trained on the CICIDS2017 benchmark dataset to categorize threats into seven distinct classes — Benign, DDoS, Brute Force, Malware, Data Exfiltration, Reconnaissance, and SQL Injection.')
    add_body(doc, "The system features a comprehensive modular architecture encompassing format-agnostic log ingestion supporting JSON, CSV, Syslog, and Windows EVTX formats; a multi-stage preprocessing pipeline for data normalization and validation; an advanced feature engineering module extracting 78 engineered features across statistical, temporal, behavioral, and network dimensions; a weighted ensemble coordinator that synthesizes anomaly scores and classification outputs into prioritized alerts (P0–P4) with MITRE ATT&CK framework mapping; and an explainability layer generating human-readable reasoning chains, feature importance rankings, and contextual response recommendations aligned with the NIST Incident Response Framework.")
    add_body(doc, "SOC Copilot operates as a fully offline, desktop-based application built with PyQt6, ensuring complete data sovereignty and compliance with air-gapped security environments. The system employs a governance-first design philosophy with analyst-in-the-loop oversight, complete auditability through append-only audit trails, and configurable kill-switch mechanisms. Experimental evaluation demonstrates classification accuracy exceeding 99%, significant reduction in analyst workload through automated triaging and intelligent alert deduplication, and faster incident response through contextual alert enrichment and actionable recommendations. The project establishes a foundation for next-generation SOC automation that augments human analytical capabilities rather than replacing them.")

# ══════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ══════════════════════════════════════════
def write_chapter1(doc):
    add_page_break(doc)
    add_heading(doc, "CHAPTER 1: INTRODUCTION", 1)

    add_heading(doc, "1.1 Background", 2)
    add_body(doc, "The digital transformation of enterprises, governments, and critical infrastructure has fundamentally reshaped the cybersecurity landscape. Organizations increasingly depend on interconnected digital systems for core operations, creating vast and complex attack surfaces that malicious actors continuously seek to exploit. According to the IBM Cost of a Data Breach Report 2023, the global average cost of a data breach reached $4.45 million, with organizations taking an average of 277 days to identify and contain a breach. The threat landscape has evolved from isolated attacks by individual hackers to orchestrated campaigns by nation-state actors, organized cybercrime syndicates, and advanced persistent threat (APT) groups employing sophisticated multi-stage attack methodologies.")
    add_body(doc, "The proliferation of cloud computing, Internet of Things (IoT) devices, remote work infrastructures, and Software-as-a-Service (SaaS) applications has exponentially increased the volume and diversity of security-relevant data generated across organizational networks. Modern enterprises generate millions of log events daily from firewalls, intrusion detection and prevention systems (IDS/IPS), endpoint detection and response (EDR) agents, application servers, authentication systems, and cloud workload monitors. This data deluge presents both an opportunity and a challenge: while comprehensive logging provides the raw material for threat detection, the sheer volume overwhelms traditional analysis approaches.")

    add_heading(doc, "1.2 Problem Statement", 2)
    add_body(doc, "Security Operations Centers serve as the centralized nerve center for organizational cybersecurity, responsible for continuous monitoring, threat detection, incident investigation, and response coordination. However, modern SOC workflows face several systemic challenges that collectively undermine operational effectiveness:")
    add_bullet(doc, "Alert Overload and Fatigue: Enterprise SOCs routinely generate tens of thousands of alerts per day, with studies indicating that up to 95% of these alerts are false positives. This phenomenon desensitizes analysts and increases the probability that genuine threats are overlooked.")
    add_bullet(doc, "Slow Response Times: Manual log analysis and threat investigation are inherently time-intensive. Analysts must cross-reference multiple data sources, correlate indicators of compromise (IoCs), and apply domain expertise, contributing to extended Mean Time to Detect (MTTD) and Mean Time to Respond (MTTR).")
    add_bullet(doc, "Lack of Contextual Intelligence: Traditional SIEM systems excel at log aggregation but lack the contextual intelligence necessary to distinguish between benign anomalies and genuine threats.")
    add_bullet(doc, "Skills Shortage: The cybersecurity workforce gap — estimated at 3.4 million unfilled positions globally — exacerbates the burden on existing analysts.")

    add_heading(doc, "1.3 Need for AI-Powered SOC Assistant", 2)
    add_body(doc, "The limitations of manual analysis and rule-based detection have motivated significant research into the application of artificial intelligence and machine learning for cybersecurity. AI-driven security assistants can automate the analysis of vast log volumes, identify subtle threat patterns that evade rule-based detection, prioritize alerts based on comprehensive risk assessment, and generate contextually enriched response recommendations. By augmenting analyst capabilities rather than replacing human judgment, AI-powered assistants address the scalability, accuracy, and speed requirements of modern SOC operations while maintaining the governance and oversight essential for security-critical decisions.")

    add_heading(doc, "1.4 Objectives of SOC Copilot", 2)
    add_bullet(doc, "Develop a hybrid AI-powered system that combines unsupervised anomaly detection with supervised multi-class attack classification for comprehensive threat detection.")
    add_bullet(doc, "Implement format-agnostic log ingestion supporting JSON, CSV, Syslog, and Windows EVTX formats with automated parsing and normalization.")
    add_bullet(doc, "Design and extract 78 engineered features across statistical, temporal, behavioral, and network dimensions for robust threat characterization.")
    add_bullet(doc, "Create an ensemble scoring mechanism that synthesizes dual-model outputs into prioritized alerts (P0–P4) with MITRE ATT&CK framework mapping.")
    add_bullet(doc, "Build an explainability layer providing human-readable reasoning, feature importance, and contextual response recommendations.")
    add_bullet(doc, "Develop a fully offline, desktop-based application with governance-first design, analyst-in-the-loop oversight, and complete auditability.")

    add_heading(doc, "1.5 Scope of the Project", 2)
    add_body(doc, "The scope of SOC Copilot encompasses the complete security log analysis pipeline from ingestion through alert generation and visualization. The system processes batch and real-time log data from multiple format types, applies AI-driven analysis for threat detection and classification, and presents actionable intelligence through an interactive desktop dashboard. The project scope includes model training using the CICIDS2017 benchmark dataset, real-time dashboard with animated metrics, analyst feedback integration for continuous improvement, and governance controls including kill-switch mechanisms. The system is designed for deployment in air-gapped and high-security environments where data sovereignty and offline operation are mandatory requirements.")

# ══════════════════════════════════════════
# CHAPTER 2: LITERATURE REVIEW
# ══════════════════════════════════════════
def write_chapter2(doc):
    add_page_break(doc)
    add_heading(doc, "CHAPTER 2: LITERATURE REVIEW", 1)

    add_heading(doc, "2.1 Existing SIEM Systems", 2)
    add_body(doc, "Security Information and Event Management (SIEM) platforms represent the foundational infrastructure of modern Security Operations Centers. Leading commercial solutions include Splunk Enterprise Security, IBM QRadar, ArcSight (Micro Focus), and Microsoft Sentinel. These platforms provide centralized log aggregation from diverse sources, real-time event correlation using predefined rules, compliance reporting and audit trail maintenance, and alert generation based on signature matching and threshold violations. While SIEM systems have been instrumental in standardizing security monitoring workflows, their detection capabilities are fundamentally constrained by their rule-based architecture.")

    add_heading(doc, "2.2 AI in Cybersecurity", 2)
    add_body(doc, "The application of artificial intelligence to cybersecurity has seen rapid advancement across multiple paradigms. Machine learning approaches for network intrusion detection have been explored using Support Vector Machines (SVMs), Decision Trees, k-Nearest Neighbors (k-NN), and ensemble methods. Deep learning architectures including Convolutional Neural Networks (CNNs), Recurrent Neural Networks (RNNs), Long Short-Term Memory (LSTM) networks, and Autoencoders have shown promising results in detecting complex attack patterns. Recent advances in Natural Language Processing (NLP) and Large Language Models (LLMs) have introduced new possibilities for threat intelligence analysis, log summarization, and automated incident report generation.")

    add_heading(doc, "2.3 Limitations of Traditional SOC Workflows", 2)
    add_body(doc, "Traditional SOC workflows suffer from several well-documented limitations: (1) Rule Rigidity — static correlation rules cannot adapt to novel attack patterns without manual updates; (2) Alert Volume — the average enterprise SOC receives over 10,000 alerts daily, with analysts able to investigate only a fraction; (3) Context Deficit — alerts lack the behavioral and contextual information necessary for informed decision-making; (4) Inconsistency — manual analysis introduces variability across analyst shifts and skill levels; (5) Scalability Constraints — linear scaling of analyst headcount is economically unsustainable.")

    add_heading(doc, "2.4 Black-Box AI vs. Explainable AI", 2)
    add_body(doc, "A significant concern in applying AI to security operations is the interpretability of model decisions. Deep learning models, while achieving high accuracy, operate as black-box systems whose internal decision logic is opaque to human operators. In security-critical contexts, this opacity undermines analyst trust and complicates regulatory compliance. Explainable AI (XAI) techniques aim to make model predictions interpretable by providing feature importance rankings, decision path visualization, counterfactual explanations, and confidence calibration. SOC Copilot addresses explainability through a wrapper-based approach that provides human-readable reasoning chains alongside every alert, enabling analysts to understand and validate AI-generated assessments.")

    add_heading(doc, "2.5 Research Gaps", 2)
    add_body(doc, "The literature review reveals several gaps that SOC Copilot addresses: (1) Most proposed systems focus on a single detection paradigm rather than combining complementary signals through ensemble coordination; (2) Few systems integrate explainability as a first-class design requirement; (3) Governance and analyst oversight mechanisms are rarely addressed; (4) Most systems assume cloud connectivity, making them unsuitable for air-gapped environments; (5) End-to-end pipeline integration from log ingestion through alert visualization is seldom demonstrated.")

    add_heading(doc, "2.6 Comparative Analysis", 2)
    add_table(doc,
        ["Feature", "Traditional SIEM", "ML-Based IDS", "SOC Copilot"],
        [
            ["Detection Method", "Rule-based", "Single ML Model", "Hybrid Ensemble (IF + RF)"],
            ["Novel Threat Detection", "No", "Partial", "Yes (Anomaly Detection)"],
            ["Multi-class Classification", "No", "Limited", "7 Attack Classes"],
            ["Explainability", "None", "Minimal", "Full XAI Layer"],
            ["MITRE ATT&CK Mapping", "Manual", "Rare", "Automated"],
            ["Offline Operation", "Partial", "Varies", "Fully Offline"],
            ["Governance Controls", "Audit Logs", "None", "Kill Switch + Audit"],
            ["Analyst Feedback Loop", "No", "Rare", "Built-in"],
            ["Feature Engineering", "N/A", "Basic", "78 Engineered Features"],
        ],
        col_widths=[1.5, 1.5, 1.5, 1.8])


def write_chapter3(doc):
    add_page_break(doc)
    add_heading(doc, "CHAPTER 3: SYSTEM ANALYSIS", 1)

    add_heading(doc, "3.1 Introduction", 2)
    add_body(doc, "This chapter presents a comprehensive analysis of the system requirements, the existing landscape of SOC tools, and the proposed SOC Copilot system. The analysis follows a structured approach encompassing the software development lifecycle, requirements identification, evaluation of existing systems, and detailed specification of the proposed system architecture and capabilities.")

    add_heading(doc, "3.1.1 Software Development Life Cycle", 3)
    add_body(doc, "SOC Copilot was developed following the Agile-Iterative software development model, executed across 17 development sprints. This methodology was chosen for its flexibility in accommodating evolving requirements, iterative refinement based on testing feedback, and support for incremental feature delivery. Each sprint encompassed planning, design, implementation, testing, and review phases, with sprint durations ranging from 5 to 10 days. Key milestones included: Sprint 1-3 (Core Data Pipeline), Sprint 4-7 (ML Model Development), Sprint 8-10 (Intelligence Layer), Sprint 11-14 (UI and Dashboard), Sprint 15-17 (Production Readiness and Governance).")

    add_heading(doc, "3.1.2 Requirements Identification", 3)
    add_body(doc, "Requirements were identified through analysis of industry SOC workflows, review of academic literature on AI-driven threat detection, evaluation of existing SIEM system limitations, and consultation with cybersecurity domain knowledge resources including MITRE ATT&CK, NIST Cybersecurity Framework, and OWASP guidelines.")

    add_heading(doc, "3.2 Existing System", 2)
    add_heading(doc, "3.2.1 Manual Log Monitoring", 3)
    add_body(doc, "In traditional SOC environments, analysts manually review log entries from multiple sources using command-line tools, log viewers, and spreadsheet applications. This approach requires significant domain expertise, is inherently slow, and cannot scale to the volume of logs generated by modern enterprise networks.")
    add_heading(doc, "3.2.2 Traditional SIEM Systems", 3)
    add_body(doc, "Commercial SIEM platforms such as Splunk, IBM QRadar, and ArcSight provide centralized log aggregation, correlation rule engines, and compliance dashboards. While these systems automate log collection and basic alerting, they rely on manually maintained rule sets, generate excessive false positives, lack contextual intelligence, and require significant infrastructure investment.")
    add_heading(doc, "3.2.3 Limitations of Existing Systems", 3)
    add_bullet(doc, "Static detection rules cannot adapt to novel attack patterns without manual updates.")
    add_bullet(doc, "High false positive rates (up to 95%) lead to analyst desensitization and alert fatigue.")
    add_bullet(doc, "Lack of automated threat classification and risk prioritization.")
    add_bullet(doc, "No built-in explainability for detection decisions.")
    add_bullet(doc, "Cloud-dependent architectures unsuitable for air-gapped environments.")
    add_bullet(doc, "Limited or no integration of machine learning for behavioral analysis.")

    add_heading(doc, "3.3 Proposed System", 2)
    add_body(doc, "SOC Copilot addresses the limitations of existing systems through a comprehensive AI-powered architecture comprising the following core components:")
    add_heading(doc, "AI Log Analysis Engine", 3)
    add_body(doc, "The analysis engine employs a hybrid ensemble of Isolation Forest (unsupervised anomaly detection) and Random Forest (supervised multi-class classification). Isolation Forest establishes behavioral baselines from normal network activity and assigns anomaly scores normalized to [0, 1], enabling detection of novel threats without labeled training data. Random Forest classifies log records into seven attack categories with associated confidence scores.")
    add_heading(doc, "Threat Classification", 3)
    add_body(doc, "The system classifies threats into seven distinct categories: Benign, DDoS, Brute Force, Malware, Data Exfiltration, Reconnaissance, and SQL Injection. Each classification includes confidence scores, probability distributions across all classes, and mapping to relevant MITRE ATT&CK tactics and techniques.")
    add_heading(doc, "Risk Scoring", 3)
    add_body(doc, "The ensemble coordinator computes a combined risk score using the formula: Combined Risk Score = 0.4 × Anomaly Score + 0.6 × (Threat Severity × Classification Confidence). This weighted combination ensures that both anomalous behavior and identified attack patterns contribute to the overall risk assessment.")
    add_heading(doc, "Incident Prioritization", 3)
    add_body(doc, "Alerts are assigned priority levels from P0 (Critical) through P4 (Informational) based on the combined risk score and threat classification. P0 and P1 alerts require immediate analyst attention, while P3 and P4 alerts are logged for trend analysis.")
    add_heading(doc, "Automated Alerting", 3)
    add_body(doc, "The alert generation module creates structured alert objects containing timestamp, source details, threat classification, risk score, priority level, MITRE ATT&CK mapping, feature importance rankings, and recommended response actions. An event deduplication mechanism prevents alert flooding from repetitive benign events.")
    add_heading(doc, "Explainable AI Outputs", 3)
    add_body(doc, "Every alert generated by SOC Copilot includes human-readable explanations covering: why the event was flagged, which features contributed most to the detection, the confidence level of the classification, relevant MITRE ATT&CK tactics, and specific recommended actions aligned with the NIST Incident Response Framework.")


def write_chapter4(doc):
    add_page_break(doc)
    add_heading(doc, "CHAPTER 4: SYSTEM DESIGN", 1)

    add_heading(doc, "4.1 Overall Architecture", 2)
    add_body(doc, "SOC Copilot employs a layered, phased architecture designed for modularity, extensibility, and governance compliance. The architecture comprises five principal layers:")
    add_image(doc, "system_architecture.png", "Figure 4.1: SOC Copilot — System Architecture Diagram")

    add_heading(doc, "Frontend (PyQt6 Desktop UI)", 3)
    add_body(doc, "The presentation layer implements a PyQt6-based desktop application featuring a sidebar navigation system with keyboard shortcuts, a system status bar with LED-style indicators for pipeline health monitoring, and a stacked page layout supporting Dashboard, Alerts, Investigation, AI Assistant, and Settings views. The interface communicates with the backend through a read-only Controller Bridge pattern ensuring unidirectional data flow and preventing the UI from modifying core analysis state.")

    add_heading(doc, "Backend (Application Controller)", 3)
    add_body(doc, "The Application Controller serves as the central orchestration layer, managing the end-to-end processing pipeline including log ingestion, preprocessing, feature extraction, model inference, and alert generation. The controller implements a micro-batching architecture for real-time processing with configurable batch sizes and processing intervals. It exposes a clean API through the Controller Bridge for UI consumption.")

    add_heading(doc, "Database (SQLite)", 3)
    add_body(doc, "SOC Copilot uses SQLite databases for persistent storage: a Result Store for analysis outputs, a Feedback Store for analyst verdicts enabling post-hoc analysis of model accuracy, a Governance Database providing append-only audit trails, and a Drift Monitoring store for feature distribution snapshots used to detect model degradation over time.")

    add_heading(doc, "AI/ML Models", 3)
    add_body(doc, "Pre-trained model artifacts are serialized via joblib and loaded in read-only mode at application startup. The Isolation Forest model is trained on benign network traffic from the CICIDS2017 dataset. The Random Forest classifier is trained on the full CICIDS2017 dataset with balanced class weights to handle class imbalance across seven attack categories.")

    add_heading(doc, "Log Ingestion Pipeline", 3)
    add_body(doc, "The ingestion pipeline uses a Parser Factory pattern supporting JSON, CSV, Syslog, and EVTX formats. Each parser implements format-specific extraction logic and produces standardized ParsedRecord objects. The pipeline includes schema validation, malformed record handling, and support for both batch file processing and directory-level ingestion.")

    add_heading(doc, "Visualization Dashboard", 3)
    add_body(doc, "The dashboard provides real-time threat monitoring through animated metric cards displaying total events processed, threats detected, critical alerts, and average risk score. The alerts view presents a priority-sorted, filterable table with color-coded severity indicators and incremental real-time updates via 2-second polling intervals.")

    add_heading(doc, "4.2 Input Design", 2)
    add_heading(doc, "Log Ingestion", 3)
    add_body(doc, "The system accepts security log files through a file upload interface in the desktop application. Supported formats include JSON/JSONL (structured log events from cloud services and APIs), CSV (tabular log exports from legacy systems and SIEM exports), Syslog (RFC 5424/3164 formatted messages from network devices), and Windows EVTX (binary event log format from Windows operating systems).")
    add_heading(doc, "JSON Parsing", 3)
    add_body(doc, "JSON and JSONL log files are parsed using Python's built-in json module with support for both single JSON objects and line-delimited JSON streams. The parser extracts standard fields including timestamp, source IP, destination IP, port, protocol, event type, and severity, with automatic field name mapping for common naming conventions.")
    add_heading(doc, "Threat Intelligence Feeds", 3)
    add_body(doc, "The system incorporates static MITRE ATT&CK mapping data for threat context enrichment. Each detected threat class is mapped to relevant ATT&CK tactics and techniques, providing analysts with standardized threat characterization. Future versions will support dynamic threat intelligence feed integration.")

    add_heading(doc, "4.3 Output Design", 2)
    add_heading(doc, "Threat Classification Output", 3)
    add_body(doc, "Each analyzed log record produces a classification result containing the predicted attack class, confidence score, and probability distribution across all seven classes. The classification output enables analysts to assess not only the most likely threat category but also alternative classifications.")
    add_heading(doc, "Risk Score", 3)
    add_body(doc, "The combined risk score (0.0–1.0) synthesizes anomaly detection and classification outputs, providing a single metric for alert prioritization. Risk scores are mapped to four risk levels: Critical (≥0.8), High (≥0.6), Medium (≥0.4), and Low (<0.4).")
    add_heading(doc, "Incident Summary", 3)
    add_body(doc, "Each alert includes a structured incident summary with threat description, affected assets, recommended actions, and relevant MITRE ATT&CK context. The summary is designed for rapid analyst consumption and supports direct inclusion in incident response documentation.")
    add_heading(doc, "Visual Dashboards", 3)
    add_body(doc, "The dashboard presents threat metrics through animated counter cards, priority distribution charts, risk level gauges, and timeline visualizations. Color-coded alert tables provide at-a-glance severity assessment.")

    add_heading(doc, "4.4 Data Flow Diagram", 2)
    add_body(doc, "The system follows a strictly sequential data flow: Log Files → Format Detection → Parsing → Validation → Normalization → Feature Extraction (78 features) → Isolation Forest Scoring → Random Forest Classification → Ensemble Coordination → Risk Score Calculation → Alert Generation → Explainability → Dashboard Display → Analyst Review → Feedback Storage. Phase isolation ensures downstream phases do not modify the upstream Detection Engine behavior, maintaining architectural integrity and auditability.")
    add_image(doc, "data_flow_diagram.png", "Figure 4.2: SOC Copilot — Data Flow Diagram (Level 1)")
    add_body(doc, "Level 0 (Context Diagram): External entities include Log Sources, SOC Analyst, Configuration Files, and MITRE ATT&CK Framework. The SOC Copilot system receives raw logs and configuration inputs and produces alerts, risk scores, and dashboard metrics as outputs.")
    add_body(doc, "Level 1 (Major Processes): The system decomposes into eight major processes: (1) Log Ingestion & Parsing, (2) Preprocessing & Normalization, (3) Feature Engineering, (4) ML Inference, (5) Ensemble Scoring, (6) Alert Generation, (7) Event Deduplication, and (8) Dashboard & UI. Data stores include the Result Store, Feedback Store (SQLite), and Audit Log.")

    add_heading(doc, "4.5 UML Diagrams", 2)

    add_heading(doc, "4.5.1 Use Case Diagram", 3)
    add_body(doc, "The Use Case Diagram identifies three primary actors: SOC Analyst (uploads logs, views dashboard, investigates alerts, provides feedback), System/Scheduler (automates ingestion, parsing, feature extraction, ML inference, scoring, alert generation, and drift monitoring), and Administrator (configures thresholds, manages kill switch, exports/imports feedback data). Key use cases include: Upload Log Files, Ingest & Parse Logs, Run ML Inference, Generate MITRE-Mapped Alerts, View Real-Time Dashboard, Investigate Alerts with AI Explanation, and Configure System Parameters.")
    add_image(doc, "use_case_diagram.png", "Figure 4.3: SOC Copilot — Use Case Diagram")

    add_heading(doc, "4.5.2 Class Diagram", 3)
    add_body(doc, "The Class Diagram depicts the object-oriented architecture of SOC Copilot with the following key classes: BaseParser (abstract) with concrete implementations JSONParser, CSVParser, SyslogParser, and EVTXParser; BaseDetector and BaseClassifier abstractions for ML models; ParserFactory using the Factory pattern; PreprocessingPipeline and FeatureEngineeringPipeline for data transformation; ModelInference, EnsembleCoordinator, AlertGenerator, and EventDeduplicator for the analysis pipeline; SOCCopilot as the main pipeline orchestrator; AppController and ControllerBridge for application control; and MainWindow as the UI entry point. Key data models include ParsedRecord, AnomalyResult, ClassificationResult, EnsembleResult, and Alert, with enumerations for AlertPriority (P0–P4), RiskLevel, ThreatCategory, AttackClass, and AlertStatus.")
    add_image(doc, "class_diagram.png", "Figure 4.4: SOC Copilot — Class Diagram")

    add_heading(doc, "4.5.3 Sequence Diagram", 3)
    add_body(doc, "The Sequence Diagram illustrates the end-to-end log analysis flow: (1) SOC Analyst uploads a log file through the UI; (2) MainWindow forwards the request to ControllerBridge; (3) ControllerBridge invokes AppController.process_batch(); (4) AppController delegates to SOCCopilot Pipeline; (5) ParserFactory parses the file into ParsedRecords; (6) PreprocessingPipeline transforms records into a clean DataFrame; (7) FeatureEngineeringPipeline extracts feature vectors; (8) AnalysisPipeline iterates over records invoking ModelInference for anomaly scoring and classification; (9) EnsembleCoordinator computes combined scores; (10) AlertGenerator creates alerts for threat events while EventDeduplicator suppresses benign duplicates; (11) Results propagate back through the Controller Bridge to update the Dashboard.")
    add_image(doc, "sequence_diagram.png", "Figure 4.5: SOC Copilot — Sequence Diagram")

    add_heading(doc, "4.5.4 Activity Diagram", 3)
    add_body(doc, "The Activity Diagram models the workflow from log upload to alert display: Start → Upload Log File → Detect Format → Parse Records → Validate Schema → [Valid?] → Normalize Fields → Extract Features → Run Anomaly Detection → Run Classification → Compute Ensemble Score → [Requires Alert?] → Generate Alert with MITRE Mapping → Deduplicate → Store Result → Update Dashboard → End. Invalid records branch to a Log Error activity before rejoining the main flow. Non-alertable events branch to Store Benign Record before reaching the End state.")
    add_image(doc, "activity_diagram.png", "Figure 4.6: SOC Copilot — Activity Diagram")


if __name__ == "__main__":
    print("Part 1 loaded. Use generate_project_report_main.py to generate the full report.")
