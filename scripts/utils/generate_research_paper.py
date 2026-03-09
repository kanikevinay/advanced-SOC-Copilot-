"""
Generate IJIRT Research Paper matching the official IJIRT paper template format.

Template format specification (extracted from ijirt_paperformat.docx):
  - Page: US Letter (12240 x 15840 twips)
  - Margins: 1 inch all sides (1440 twips)
  - Section 0: Single column (title block)
  - Section 1: Two-column layout, column gap 720 twips (0.5 inch)
  - Font: Times New Roman 10pt body throughout
  - Title: 24pt bold centered
  - Authors: 11pt centered
  - Section Headings: 10pt centered, space before/after 6pt
  - Sub-section Headings: 10pt italic justified
  - Body: 10pt justified, space after 0
  - Header: IJIRT banner with ISSN
  - IEEE-style references
"""

import os
import copy
from docx import Document
from docx.shared import Pt, Inches, Emu, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────
# Constants matching IJIRT template exactly
# ──────────────────────────────────────────────
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(10)       # 127000 EMU = 10pt
TITLE_SIZE = Pt(24)      # 304800 EMU = 24pt
AUTHOR_SIZE = Pt(11)
HEADING_SPACE = Pt(6)    # 76200 EMU = 6pt
REF_SIZE = Pt(9)
HEADER_FONT = "Arial"
HEADER_SIZE = Pt(13)     # 165100 EMU = 13pt


def make_section_break(doc, col_num=None, col_space=720, break_type="continuous"):
    """Insert a section break with optional column settings."""
    new_section = doc.add_section()
    sect_pr = new_section._sectPr

    # Page size: Letter
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        sect_pr.append(pg_sz)
    pg_sz.set(qn("w:w"), "12240")
    pg_sz.set(qn("w:h"), "15840")

    # Margins: 1 inch all sides
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        sect_pr.append(pg_mar)
    for attr in ["w:top", "w:right", "w:bottom", "w:left"]:
        pg_mar.set(qn(attr), "1440")
    pg_mar.set(qn("w:header"), "720")
    pg_mar.set(qn("w:footer"), "720")
    pg_mar.set(qn("w:gutter"), "0")

    # Section type
    if break_type:
        type_elem = sect_pr.find(qn("w:type"))
        if type_elem is None:
            type_elem = OxmlElement("w:type")
            sect_pr.insert(0, type_elem)
        type_elem.set(qn("w:val"), break_type)

    # Columns
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    if col_num and col_num > 1:
        cols.set(qn("w:num"), str(col_num))
    cols.set(qn("w:space"), str(col_space))

    return new_section


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def add_formatted_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document matching academic style."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.font.name = FONT_NAME

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    # Small spacing after table
    sp = doc.add_paragraph("")
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)
    return table


def create_paper():
    doc = Document()

    # ─── Configure first section (single column for title) ───
    section0 = doc.sections[0]
    pg_sz = section0._sectPr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        section0._sectPr.append(pg_sz)
    pg_sz.set(qn("w:w"), "12240")
    pg_sz.set(qn("w:h"), "15840")

    section0.top_margin = Twips(1440)
    section0.bottom_margin = Twips(1440)
    section0.left_margin = Twips(1440)
    section0.right_margin = Twips(1440)

    cols0 = section0._sectPr.find(qn("w:cols"))
    if cols0 is None:
        cols0 = OxmlElement("w:cols")
        section0._sectPr.append(cols0)
    cols0.set(qn("w:space"), "720")

    # ─── Setup header (IJIRT banner) ───
    header = section0.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    parts = [
        ("\u00a9 2026", False),
        (" | ", False),
        ("IJIRT", True),
        (" | ", False),
        ("Volume 12 Issue 9", False),
        (" | ", False),
        ("ISSN: 2349-6002", False),
    ]
    for text, bold in parts:
        run = header_para.add_run(text)
        run.font.name = HEADER_FONT
        run.font.size = HEADER_SIZE
        run.bold = bold

    # ─── Default paragraph style ───
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ══════════════════════════════════════════════
    #  HELPER FUNCTIONS
    # ══════════════════════════════════════════════

    def add_section_heading(text):
        """Roman-numeral section headings: centered, 10pt, space before/after 6pt."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = HEADING_SPACE
        p.paragraph_format.space_after = HEADING_SPACE
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        run.bold = False
        return p

    def add_subsection_heading(text):
        """Lettered sub-section headings: italic, justified, 10pt."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        run.italic = True
        return p

    def add_body(text):
        """Body paragraph: justified, 10pt, Times New Roman, spaceAfter=0."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        return p

    def add_bold_body(bold_part, rest):
        """Body paragraph with leading bold text."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(bold_part)
        r1.bold = True
        r1.font.name = FONT_NAME
        r1.font.size = BODY_SIZE
        r2 = p.add_run(rest)
        r2.font.name = FONT_NAME
        r2.font.size = BODY_SIZE
        return p

    def add_italic_body(text):
        """Italic body paragraph."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        run.italic = True
        return p

    def add_bullet(text):
        """Bullet point in body style."""
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        # Clear default run and add properly formatted one
        p.clear()
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        return p

    # ══════════════════════════════════════════════════════
    #  TITLE BLOCK (Single Column Section)
    # ══════════════════════════════════════════════════════

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "SOC Copilot: An AI-Powered Security Operations Assistant "
        "for Automated Threat Detection and Intelligent Incident Response"
    )
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = FONT_NAME
    title.paragraph_format.space_after = Pt(6)

    # ── Authors ──
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run("Author 1, Author 2, Author 3")
    run.font.size = AUTHOR_SIZE
    run.font.name = FONT_NAME
    authors.paragraph_format.space_after = Pt(2)

    # ── Department ──
    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dept.add_run("Department of Computer Science and Engineering")
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME

    # ── Institution ──
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inst.add_run("[Institution Name], [City, State, Country]")
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME
    inst.paragraph_format.space_after = Pt(6)

    # ── Abstract heading ──
    abs_h = doc.add_paragraph()
    abs_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abs_h.paragraph_format.space_before = Pt(6)
    abs_h.paragraph_format.space_after = Pt(4)
    run = abs_h.add_run("Abstract")
    run.bold = True
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME

    # ── Abstract text ──
    abs_text = doc.add_paragraph()
    abs_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_text.paragraph_format.space_after = Pt(4)
    run = abs_text.add_run(
        "The proliferation of sophisticated cyber threats has placed unprecedented pressure on "
        "Security Operations Centers (SOCs), where analysts are tasked with monitoring, triaging, "
        "and responding to thousands of security alerts daily. This overwhelming volume of "
        "heterogeneous log data, compounded by high false-positive rates and limited contextual "
        "intelligence, leads to a well-documented phenomenon known as alert fatigue \u2014 a critical "
        "bottleneck that degrades threat detection efficacy and increases mean time to detect (MTTD) "
        "and mean time to respond (MTTR). Traditional Security Information and Event Management "
        "(SIEM) systems, while capable of log aggregation, rely predominantly on static rule-based "
        "detection mechanisms that fail to adapt to evolving attack vectors and generate excessive "
        "noise. This paper presents SOC Copilot, an AI-powered security operations assistant that "
        "leverages a hybrid machine learning ensemble \u2014 combining Isolation Forest for unsupervised "
        "anomaly detection with Random Forest for supervised multi-class attack classification \u2014 to "
        "automate threat detection, prioritize alerts, and generate intelligent, explainable "
        "incident response recommendations. The system ingests security logs from multiple formats "
        "(JSON, CSV, Syslog, Windows EVTX), extracts 78 engineered features across statistical, "
        "temporal, behavioral, and network dimensions, and produces prioritized alerts (P0\u2013P4) with "
        "MITRE ATT&CK mapping and human-readable reasoning. SOC Copilot operates entirely offline "
        "with a governance-first design philosophy, ensuring analyst-in-the-loop oversight, complete "
        "auditability, and data sovereignty compliance. Experimental evaluation demonstrates a "
        "classification accuracy exceeding 99%, significant reduction in analyst workload through "
        "automated triaging, and faster incident response through contextual alert enrichment and "
        "actionable recommendations."
    )
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE
    run.italic = True

    # ── Keywords ──
    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw.paragraph_format.space_after = Pt(6)
    r1 = kw.add_run("Keywords \u2014 ")
    r1.bold = True
    r1.font.name = FONT_NAME
    r1.font.size = BODY_SIZE
    r2 = kw.add_run(
        "Cybersecurity; SOC Automation; Threat Detection; Log Analysis; "
        "AI Security Assistant; Machine Learning in Security; Incident Response; SIEM Enhancement"
    )
    r2.font.name = FONT_NAME
    r2.font.size = BODY_SIZE
    r2.italic = True

    # ══════════════════════════════════════════════════════
    #  SWITCH TO TWO-COLUMN LAYOUT
    # ══════════════════════════════════════════════════════
    make_section_break(doc, col_num=2, col_space=720)

    # Also set the header for the new section
    section1 = doc.sections[-1]
    header1 = section1.header
    header1.is_linked_to_previous = True

    # ══════════════════════════════════════════════════════
    #  I. INTRODUCTION
    # ══════════════════════════════════════════════════════
    add_section_heading("I. INTRODUCTION")

    add_body(
        "The digital transformation of enterprises, governments, and critical infrastructure has "
        "expanded the attack surface available to malicious actors at an unprecedented rate. "
        "According to the IBM Cost of a Data Breach Report, the global average cost of a data "
        "breach reached $4.45 million in 2023, with organizations taking an average of 277 days "
        "to identify and contain a breach [1]. This escalating threat landscape \u2014 characterized by "
        "advanced persistent threats (APTs), zero-day exploits, ransomware campaigns, and supply "
        "chain attacks \u2014 demands robust, intelligent, and responsive security operations."
    )

    add_body(
        "Security Operations Centers (SOCs) serve as the nerve center of organizational "
        "cybersecurity, responsible for continuous monitoring, threat detection, incident "
        "investigation, and response coordination. SOC analysts rely on Security Information and "
        "Event Management (SIEM) systems to aggregate and correlate security logs from diverse "
        "sources including firewalls, intrusion detection/prevention systems (IDS/IPS), endpoint "
        "detection and response (EDR) agents, application servers, and network devices. However, "
        "modern SOC workflows face several systemic challenges that impede operational effectiveness."
    )

    add_bold_body(
        "Alert Overload and Fatigue. ",
        "Enterprise SOCs routinely generate tens of thousands of alerts per day, with studies "
        "indicating that up to 95% of these alerts are false positives [2]. This phenomenon, "
        "known as alert fatigue, desensitizes analysts and increases the probability that genuine "
        "threats are overlooked or inadequately investigated. The cognitive burden of manually "
        "reviewing high-volume, low-fidelity alerts leads to inconsistent decision-making across "
        "analyst shifts."
    )

    add_bold_body(
        "Slow Response Times. ",
        "Manual log analysis and threat investigation are inherently time-intensive processes. "
        "Analysts must cross-reference multiple data sources, correlate indicators of compromise "
        "(IoCs), and apply domain expertise to determine the severity and scope of potential "
        "incidents. This manual workflow contributes to extended MTTD and MTTR, providing "
        "attackers with prolonged dwell time within compromised environments."
    )

    add_bold_body(
        "Lack of Contextual Intelligence. ",
        "Traditional SIEM systems excel at log aggregation and rule-based alerting but lack the "
        "contextual intelligence necessary to distinguish between benign anomalies and genuine "
        "threats. Static detection rules cannot adapt to novel attack patterns or account for the "
        "behavioral context of network activity, resulting in both false positives and false negatives."
    )

    add_bold_body(
        "Need for AI-Driven Assistance. ",
        "The limitations of manual analysis and rule-based detection have motivated significant "
        "research into the application of artificial intelligence (AI) and machine learning (ML) "
        "for cybersecurity. AI-driven security assistants can automate the analysis of vast log "
        "volumes, identify subtle threat patterns that evade rule-based detection, prioritize "
        "alerts based on risk assessment, and generate contextually enriched response "
        "recommendations \u2014 thereby augmenting analyst capabilities rather than replacing human judgment."
    )

    add_bold_body(
        "SOC Copilot Overview. ",
        "This paper introduces SOC Copilot, an AI-powered security operations assistant that "
        "addresses the aforementioned challenges through a hybrid ML ensemble approach. SOC Copilot "
        "combines unsupervised anomaly detection (Isolation Forest) with supervised multi-class "
        "classification (Random Forest) to provide robust threat detection across both known and "
        "unknown attack vectors. The system features a modular, phased architecture encompassing "
        "log ingestion and preprocessing, ML-driven threat analysis, explainable alert generation "
        "with MITRE ATT&CK mapping, analyst feedback integration, model drift monitoring, and a "
        "governance-first control infrastructure."
    )

    add_body(
        "The remainder of this paper is organized as follows: Section II reviews related "
        "literature; Section III defines the problem statement; Section IV describes the proposed "
        "system; Section V presents the system architecture; Section VI details the methodology; "
        "Section VII discusses results; Sections VIII and IX outline advantages and applications; "
        "Section X explores future scope; and Section XI concludes the paper."
    )

    # ══════════════════════════════════════════════════════
    #  II. LITERATURE REVIEW
    # ══════════════════════════════════════════════════════
    add_section_heading("II. LITERATURE REVIEW")

    add_body(
        "The application of computational techniques to cybersecurity threat detection has been "
        "extensively studied across multiple paradigms, ranging from traditional rule-based "
        "systems to advanced deep learning approaches."
    )

    add_bold_body(
        "Traditional SIEM Systems. ",
        "Security Information and Event Management platforms such as Splunk, IBM QRadar, and "
        "ArcSight represent the foundational infrastructure of modern SOC operations [3]. These "
        "systems provide centralized log aggregation, real-time correlation, and compliance "
        "reporting capabilities. However, their detection mechanisms rely predominantly on "
        "predefined correlation rules and signature-based matching, which require continuous "
        "manual updates and cannot detect novel attack patterns. Zuech et al. [4] demonstrated "
        "that rule-based SIEM systems exhibit significant limitations in detecting sophisticated, "
        "multi-stage attacks that do not match existing signatures."
    )

    add_bold_body(
        "Rule-Based Detection Mechanisms. ",
        "Signature-based intrusion detection systems, including Snort and Suricata, compare "
        "network traffic against databases of known attack signatures. While effective for "
        "detecting well-characterized threats, Roesch [5] identified that these systems are "
        "inherently reactive \u2014 they can only detect attacks for which signatures have been "
        "previously developed. This fundamental limitation creates a detection gap for zero-day "
        "exploits, polymorphic malware, and novel attack techniques."
    )

    add_bold_body(
        "Machine Learning-Based Anomaly Detection. ",
        "The application of machine learning to network intrusion detection has garnered "
        "substantial research attention. Liu et al. [6] introduced the Isolation Forest algorithm, "
        "which detects anomalies by isolating observations through random partitioning. This "
        "unsupervised approach is particularly valuable for cybersecurity applications because it "
        "does not require labeled attack data for training and can identify novel threats based on "
        "behavioral deviation from established baselines."
    )

    add_bold_body(
        "Supervised Classification Approaches. ",
        "Random Forest classifiers, as proposed by Breiman [7], have demonstrated strong "
        "performance in multi-class attack classification tasks. The CICIDS2017 benchmark dataset "
        "evaluation by Sharafaldin et al. [8] showed that ensemble tree-based methods achieve "
        "superior accuracy in distinguishing between benign traffic and multiple attack categories "
        "including brute force, web attacks, infiltration, and botnet activity."
    )

    add_bold_body(
        "AI and Deep Learning in Cybersecurity. ",
        "Recent advances in deep learning have introduced autoencoders, recurrent neural networks "
        "(RNNs), and transformer architectures for threat detection. Shone et al. [9] proposed a "
        "deep autoencoder framework for network intrusion detection that achieved high accuracy on "
        "benchmark datasets. However, deep learning approaches often suffer from interpretability "
        "limitations \u2014 a significant concern in security contexts where analysts must understand "
        "and validate detection decisions [10]."
    )

    add_bold_body(
        "Research Gap. ",
        "Despite significant advances, existing research exhibits several gaps that SOC Copilot "
        "addresses. Most proposed systems focus on a single detection paradigm rather than "
        "combining complementary signals through ensemble coordination. Few systems integrate "
        "explainability as a first-class design requirement. Governance and analyst oversight "
        "mechanisms are rarely addressed. The majority of proposed systems assume cloud connectivity, "
        "making them unsuitable for air-gapped environments. SOC Copilot addresses these gaps "
        "through its hybrid ensemble approach, built-in explainability layer, governance-first "
        "architecture, and fully offline operational capability."
    )

    # ══════════════════════════════════════════════════════
    #  III. PROBLEM STATEMENT
    # ══════════════════════════════════════════════════════
    add_section_heading("III. PROBLEM STATEMENT")

    add_body(
        "Modern Security Operations Centers face a convergence of challenges that collectively "
        "undermine their ability to detect and respond to cyber threats effectively."
    )

    add_bold_body(
        "Log Volume Explosion. ",
        "Enterprise environments generate massive volumes of security-relevant log data from "
        "heterogeneous sources \u2014 firewalls, IDS/IPS sensors, endpoint agents, authentication "
        "systems, cloud workloads, and application servers. A mid-sized enterprise may produce "
        "millions of log events daily, each requiring parsing, normalization, and contextual "
        "analysis."
    )

    add_bold_body(
        "False Positive Proliferation. ",
        "Rule-based detection systems generate a disproportionate number of false positive alerts, "
        "with industry estimates suggesting that 95% or more of SIEM alerts do not correspond to "
        "actual security incidents [2]. This noise-to-signal ratio forces analysts to spend the "
        "majority of their time dismissing benign alerts rather than investigating genuine threats."
    )

    add_bold_body(
        "Analyst Fatigue and Skill Shortage. ",
        "The combination of high alert volumes, repetitive triage tasks, and the cognitive demand "
        "of complex threat investigation contributes to analyst burnout. The cybersecurity "
        "workforce gap \u2014 estimated at 3.4 million unfilled positions globally \u2014 exacerbates "
        "this challenge."
    )

    add_bold_body(
        "Lack of Contextual Intelligence. ",
        "Traditional SIEM alerts provide limited context regarding the nature, severity, and "
        "potential impact of detected events. Analysts must manually cross-reference multiple "
        "data sources, apply threat intelligence, and leverage domain expertise to contextualize alerts."
    )

    add_bold_body(
        "Manual Threat Investigation Inefficiencies. ",
        "The end-to-end threat investigation workflow \u2014 from initial alert triage through evidence "
        "collection, impact assessment, and response recommendation \u2014 requires significant manual "
        "effort. These manual processes create bottlenecks that extend incident response timelines "
        "and increase organizational risk exposure."
    )

    add_italic_body(
        "The problem can be formally stated as: Given heterogeneous multi-format security log "
        "data generated at scale, how can an intelligent system automatically parse, analyze, "
        "classify, and prioritize potential threats while providing explainable, contextual, and "
        "actionable insights that augment SOC analyst decision-making \u2014 without compromising "
        "governance oversight, data sovereignty, or operational transparency?"
    )

    # ══════════════════════════════════════════════════════
    #  IV. PROPOSED SYSTEM
    # ══════════════════════════════════════════════════════
    add_section_heading("IV. PROPOSED SYSTEM \u2013 SOC COPILOT")

    add_body(
        "SOC Copilot is designed as a comprehensive, modular security operations assistant that "
        "addresses each dimension of the problem statement through a layered, phased architecture."
    )

    add_subsection_heading("A. Log Ingestion")
    add_body(
        "The log ingestion subsystem provides format-agnostic parsing support for four widely "
        "used security log formats: JSON/JSONL (structured log events from cloud services and "
        "APIs), CSV (tabular log exports from legacy systems), Syslog (standard format used by "
        "network devices and firewalls), and Windows EVTX (binary event log format from Windows "
        "operating systems). Each parser implements format-specific extraction logic to identify "
        "and normalize critical fields including timestamps, source/destination IP addresses, "
        "ports, protocols, event types, and severity indicators."
    )

    add_subsection_heading("B. Log Parsing and Preprocessing")
    add_body(
        "Raw log records undergo a multi-stage preprocessing pipeline: (1) Timestamp "
        "Normalization \u2014 all timestamps are converted to UTC ISO 8601 format; (2) Field "
        "Standardization \u2014 source-specific field names are mapped to a canonical schema; "
        "(3) Categorical Encoding \u2014 non-numeric fields are encoded as integer representations; "
        "(4) Data Validation \u2014 records are validated against expected schemas with malformed "
        "entries logged and excluded from analysis."
    )

    add_subsection_heading("C. AI/ML-Based Analysis Engine")
    add_body(
        "The core analysis engine employs a hybrid ensemble approach combining two complementary "
        "machine learning models. Isolation Forest (unsupervised anomaly detection) is trained "
        "exclusively on benign network traffic to establish behavioral baselines, assigning "
        "anomaly scores normalized to [0, 1]. Random Forest (supervised multi-class classification) "
        "is trained on the CICIDS2017 benchmark dataset with balanced class weights to classify "
        "log records into seven categories: Benign, DDoS, Brute Force, Malware, Data Exfiltration, "
        "Reconnaissance, and SQL Injection."
    )

    add_subsection_heading("D. Threat Classification and Ensemble Logic")
    add_body(
        "The Ensemble Coordinator combines outputs from both models using a weighted decision "
        "matrix. The combined risk score is computed as: Combined Risk Score = 0.4 \u00d7 Anomaly Score "
        "+ 0.6 \u00d7 (Threat Severity \u00d7 Classification Confidence). This combined score determines the "
        "alert priority assignment (P0\u2013Critical through P4\u2013Informational) and risk level "
        "classification (Critical, High, Medium, Low). The dual-signal approach ensures that both "
        "known threats (detected by the classifier) and unknown threats (detected by the anomaly "
        "detector) are appropriately captured and prioritized."
    )

    add_subsection_heading("E. Explainability and Recommendation Engine")
    add_body(
        "SOC Copilot generates human-readable explanations for each alert through a wrapper-based "
        "explainability layer providing: classification reasoning with confidence levels, anomaly "
        "score interpretation with behavioral context, top contributing features ranked by "
        "importance, threshold crossing analysis, MITRE ATT&CK tactic and technique mapping, and "
        "suggested response actions based on threat classification. Response recommendations are "
        "aligned with NIST incident response framework categories including containment, "
        "investigation, remediation, and recovery."
    )

    add_subsection_heading("F. Dashboard Visualization")
    add_body(
        "The desktop-based user interface, implemented using PyQt6, provides a real-time "
        "monitoring dashboard with: a threat overview dashboard with animated metric cards; a "
        "priority-sorted, filterable alert table with color-coded severity indicators and "
        "incremental real-time updates (2-second polling); an investigation panel with detailed "
        "alert analysis, feature importance visualization, and reasoning breakdown; a system "
        "status bar with LED-style indicators for pipeline health; and a configuration panel for "
        "threshold calibration and governance settings."
    )

    add_subsection_heading("G. Technology Stack")

    add_formatted_table(
        doc,
        ["Component", "Technology"],
        [
            ["Core Language", "Python 3.10+"],
            ["ML Framework", "Scikit-learn (Isolation Forest, Random Forest)"],
            ["Desktop UI", "PyQt6"],
            ["Data Processing", "Pandas, NumPy"],
            ["Database", "SQLite"],
            ["Configuration", "YAML"],
            ["Testing", "Pytest (208+ tests)"],
            ["Packaging", "PyInstaller"],
            ["Version Control", "Git"],
        ],
        col_widths=[1.5, 2.0],
    )

    # ══════════════════════════════════════════════════════
    #  V. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════
    add_section_heading("V. SYSTEM ARCHITECTURE")

    add_body(
        "SOC Copilot employs a layered, phased architecture designed for modularity, "
        "extensibility, and governance compliance. The architecture comprises five principal "
        "layers operating across three functional phases."
    )

    add_subsection_heading("A. Frontend Interface Layer")
    add_body(
        "The presentation layer implements a PyQt6-based desktop application featuring a sidebar "
        "navigation system with keyboard shortcuts, a system status bar with LED indicators, and "
        "a stacked page layout supporting Dashboard, Alerts, Investigation, Assistant, and "
        "Settings views. The interface communicates with the backend through a read-only "
        "Controller Bridge pattern ensuring unidirectional data flow."
    )

    add_subsection_heading("B. Backend API Layer")
    add_body(
        "The Application Controller orchestrates the end-to-end processing pipeline, managing log "
        "ingestion, preprocessing, feature extraction, model inference, and alert generation. The "
        "controller implements a micro-batching architecture for real-time processing with "
        "configurable batch sizes and processing intervals."
    )

    add_subsection_heading("C. AI Processing Engine")
    add_body(
        "The ML inference layer loads pre-trained model artifacts (serialized via joblib) in "
        "read-only mode at application startup. Feature engineering extracts 78 numeric features "
        "across four categories: Statistical (22 features including packet counts, byte volumes, "
        "flow duration), Temporal (18 features including hour-of-day encoding), Behavioral "
        "(20 features including request frequency, error ratios), and Network (18 features "
        "including port entropy, unique destination counts)."
    )

    add_subsection_heading("D. Database Layer")
    add_body(
        "SOC Copilot uses SQLite databases for persistent storage: a Feedback Store for analyst "
        "verdicts enabling post-hoc analysis of model accuracy; a Governance Database providing "
        "an append-only audit trail; and Drift Monitoring for feature distribution snapshots."
    )

    add_subsection_heading("E. Response Generation Layer")
    add_body(
        "The Explainability and Response layer operates post-inference, generating structured "
        "alert objects containing priority level (P0\u2013P4) with MITRE ATT&CK mapping, risk "
        "classification, feature importance rankings, human-readable reasoning chains, and "
        "contextual response recommendations."
    )

    add_subsection_heading("F. End-to-End Data Flow")
    add_body(
        "The pipeline follows a strictly sequential flow: Log Files \u2192 Format Detection \u2192 "
        "Parsing \u2192 Validation \u2192 Normalization \u2192 Feature Extraction (78 features) \u2192 "
        "Isolation Forest Scoring \u2192 Random Forest Classification \u2192 Ensemble Coordination \u2192 "
        "Risk Score Calculation \u2192 Alert Generation \u2192 Explainability \u2192 Dashboard Display \u2192 "
        "Analyst Review \u2192 Feedback Storage. Phase isolation ensures downstream phases do not "
        "modify the upstream Detection Engine behavior."
    )

    # ══════════════════════════════════════════════════════
    #  VI. METHODOLOGY
    # ══════════════════════════════════════════════════════
    add_section_heading("VI. METHODOLOGY")

    add_subsection_heading("A. Log Data Processing")
    add_body(
        "The data processing pipeline implements a sequential transformation workflow: "
        "(1) Format Detection through file extension and header inspection; "
        "(2) Record Parsing via format-specific parsers; "
        "(3) Schema Validation with malformed record exclusion; "
        "(4) Timestamp Normalization to UTC ISO 8601; "
        "(5) Field Mapping to canonical schema; "
        "(6) Categorical Encoding using consistent mapping dictionaries."
    )

    add_subsection_heading("B. Feature Extraction")
    add_body(
        "The feature engineering module extracts 78 numeric features organized into four "
        "categories: Statistical (22 features \u2014 flow-level statistics including packet count, "
        "byte volume, duration, and inter-arrival time metrics), Temporal (18 features \u2014 "
        "time-based patterns including hour-of-day encoding and time-delta calculations), "
        "Behavioral (20 features \u2014 session-level patterns including request frequency, error "
        "ratios, and activity burst detection), and Network (18 features \u2014 topology metrics "
        "including port entropy and protocol distribution). Feature ordering is persisted "
        "alongside model artifacts to ensure consistent input alignment."
    )

    add_subsection_heading("C. AI Inference Workflow (Pseudo-Algorithm)")
    add_body(
        "Algorithm 1: Hybrid Ensemble Threat Detection \u2014 "
        "Input: Feature matrix X \u2208 R^(n\u00d778) from log batch. "
        "Output: Alert set A with priorities, risk scores, and explanations. "
        "Steps: (1) Load pre-trained models M_IF, M_RF; "
        "(2) For each record x_i: compute anomaly_score_i \u2190 Normalize(M_IF.decision_function(x_i)) "
        "to [0,1]; (3) Obtain (class_i, prob_i) \u2190 M_RF.predict_proba(x_i); "
        "(4) confidence_i \u2190 max(prob_i); "
        "(5) risk_score_i \u2190 0.4 \u00d7 anomaly_score_i + 0.6 \u00d7 (severity_i \u00d7 confidence_i); "
        "(6) If risk_score_i > threshold: generate explanation, MITRE mapping, recommendation; "
        "add to alert set A; (7) Return A."
    )

    add_subsection_heading("D. Threat Classification Logic")

    add_formatted_table(
        doc,
        ["Class", "Description", "Severity"],
        [
            ["Benign", "Normal network activity", "0.0"],
            ["DDoS", "Distributed denial-of-service", "0.8"],
            ["Brute Force", "Authentication attacks", "0.7"],
            ["Malware", "Malicious software execution", "1.0"],
            ["Exfiltration", "Unauthorized data transfer", "1.0"],
            ["Reconnaissance", "Network scanning/probing", "0.5"],
            ["SQL Injection", "Database exploitation", "0.9"],
        ],
        col_widths=[1.2, 1.8, 0.8],
    )

    add_subsection_heading("E. Risk Scoring")
    add_body(
        "The ensemble risk scoring combines: Anomaly Score Component (weight 0.4) capturing "
        "behavioral deviation from normal baselines, and Classification Component (weight 0.6) "
        "capturing pattern-matched threat identification. The combined score maps to four risk "
        "levels: Critical (\u2265 0.8), High (0.6\u20130.79), Medium (0.4\u20130.59), and Low (< 0.4)."
    )

    # ══════════════════════════════════════════════════════
    #  VII. RESULTS AND DISCUSSION
    # ══════════════════════════════════════════════════════
    add_section_heading("VII. RESULTS AND DISCUSSION")

    add_subsection_heading("A. System Performance")
    add_body(
        "SOC Copilot was evaluated using the CICIDS2017 benchmark dataset containing labeled "
        "network traffic across multiple attack scenarios."
    )

    add_formatted_table(
        doc,
        ["Metric", "Value"],
        [
            ["RF Classification Accuracy", "99.99%"],
            ["IF Anomaly Separation", "Confirmed"],
            ["Feature Extraction", "78 features, consistent"],
            ["Model Loading Time", "1\u20132 seconds"],
            ["Single Record Latency", "< 10 ms"],
            ["Batch (1,000 records)", "2\u20135 seconds"],
            ["Large-Scale (100K records)", "2\u20135 minutes"],
        ],
        col_widths=[2.0, 1.5],
    )

    add_subsection_heading("B. Example Log Analysis")
    add_body(
        "An illustrative analysis of a malware-associated network flow record: Input with "
        "anomalous port entropy (0.92), high unique destination count (47), and irregular "
        "inter-arrival timing. Isolation Forest Score: 0.85 (highly anomalous). Random Forest "
        "Classification: Malware (confidence: 92.5%). Combined Risk Score: 0.4 \u00d7 0.85 + 0.6 "
        "\u00d7 (1.0 \u00d7 0.925) = 0.895 \u2192 Critical (P0). Generated Reasoning: \"Classified as "
        "Malware with 92.5% confidence. High anomaly score indicates unusual behavior. Top "
        "contributing features: port_entropy, unique_destinations, time_since_last.\" "
        "Recommended Action: \"Isolate endpoint and investigate process execution.\""
    )

    add_subsection_heading("C. Comparative Evaluation")

    add_formatted_table(
        doc,
        ["Approach", "Known", "Novel", "FP Rate"],
        [
            ["Rule-Based SIEM", "High", "None", "High"],
            ["Isolation Forest Only", "Moderate", "High", "Moderate"],
            ["Random Forest Only", "High", "None", "Low"],
            ["SOC Copilot Ensemble", "High", "High", "Low"],
        ],
        col_widths=[1.5, 0.7, 0.7, 0.7],
    )

    add_subsection_heading("D. Response Time and Workload Improvement")
    add_body(
        "SOC Copilot contributes to reduced incident response timelines through automated "
        "triaging (eliminating manual alert-by-alert review), contextual enrichment (reducing "
        "initial assessment time from minutes to seconds), actionable recommendations (removing "
        "the need for independent strategy development), and MITRE ATT&CK mapping (accelerating "
        "threat intelligence correlation). The system addresses analyst workload through noise "
        "reduction via multi-model validation, priority-based focus (P0\u2013P4), batch processing "
        "of large log volumes, automated explanation generation, and a feedback loop for accuracy "
        "analysis."
    )

    # ══════════════════════════════════════════════════════
    #  VIII. ADVANTAGES
    # ══════════════════════════════════════════════════════
    add_section_heading("VIII. ADVANTAGES")

    add_body(
        "SOC Copilot offers several significant advantages: (1) Reduced SOC Workload through "
        "automated classification and prioritization; (2) Faster Triaging with sub-10ms per-record "
        "analysis latency; (3) Context-Aware Insights via feature-level reasoning and MITRE ATT&CK "
        "mapping; (4) Scalable Architecture supporting incremental capability expansion; "
        "(5) Intelligent Automation with Governance ensuring disabled-by-default controls; "
        "(6) Offline Operation and Data Sovereignty for air-gapped network compatibility; "
        "(7) Hybrid Detection covering both known signatures and novel zero-day threats; "
        "(8) Reproducibility and Auditability through deterministic scoring and append-only logging."
    )

    # ══════════════════════════════════════════════════════
    #  IX. APPLICATIONS
    # ══════════════════════════════════════════════════════
    add_section_heading("IX. APPLICATIONS")

    add_body(
        "SOC Copilot is applicable across diverse cybersecurity contexts: Enterprise SOC Teams "
        "can use it as an analyst augmentation tool; Managed Security Service Providers (MSSPs) "
        "can standardize threat detection across client environments; organizations monitoring "
        "hybrid or multi-cloud infrastructure can leverage multi-format log ingestion; Government "
        "Cybersecurity Centers benefit from fully offline architecture ensuring data sovereignty; "
        "Critical Infrastructure organizations (energy, healthcare, finance) can deploy it in "
        "restricted connectivity environments; and academic institutions can use it for "
        "cybersecurity education."
    )

    # ══════════════════════════════════════════════════════
    #  X. FUTURE SCOPE
    # ══════════════════════════════════════════════════════
    add_section_heading("X. FUTURE SCOPE")

    add_body(
        "Several enhancement trajectories are identified: (1) Real-Time SIEM Integration with "
        "streaming connectors for Splunk, Elastic SIEM, and QRadar; (2) Advanced ML Anomaly "
        "Detection through deep autoencoder integration (framework shell exists); "
        "(3) Automated Remediation via SOAR integration within the governance framework; "
        "(4) Multi-Tenant Deployment with tenant-specific configurations; "
        "(5) Cloud-Native Deployment using Docker and Kubernetes for elastic scaling; "
        "(6) Transformer-Based Sequence Analysis for multi-stage threat pattern detection; "
        "(7) LLM-Powered Natural Language Query Interface for conversational incident investigation."
    )

    # ══════════════════════════════════════════════════════
    #  XI. CONCLUSION
    # ══════════════════════════════════════════════════════
    add_section_heading("XI. CONCLUSION")

    add_body(
        "This paper presented SOC Copilot, an AI-powered security operations assistant designed "
        "to address the critical challenges facing modern Security Operations Centers. By combining "
        "Isolation Forest for unsupervised anomaly detection with Random Forest for supervised "
        "multi-class attack classification, SOC Copilot achieves robust threat detection across "
        "both known attack patterns and novel behavioral anomalies."
    )

    add_body(
        "The system's key contributions include: (1) a hybrid ML ensemble architecture with "
        "weighted decision fusion achieving classification accuracy exceeding 99%; (2) a "
        "comprehensive 78-feature engineering framework spanning statistical, temporal, behavioral, "
        "and network dimensions; (3) a built-in explainability layer providing human-readable "
        "reasoning, feature importance analysis, and MITRE ATT&CK mapping; (4) a governance-first "
        "design philosophy with disabled-by-default automation, kill switch controls, and "
        "append-only audit logging; and (5) fully offline operational capability ensuring data "
        "sovereignty and air-gap compatibility."
    )

    add_body(
        "SOC Copilot represents a significant step toward intelligent, trustworthy security "
        "automation that respects the critical role of human judgment in cybersecurity operations. "
        "As cyber threats continue to evolve in sophistication and scale, intelligent systems like "
        "SOC Copilot will become essential components of organizational cybersecurity strategy, "
        "enabling SOCs to maintain effective detection and response capabilities despite the "
        "increasing disparity between threat volume and available analyst resources."
    )

    # ══════════════════════════════════════════════════════
    #  REFERENCES
    # ══════════════════════════════════════════════════════
    add_section_heading("REFERENCES")

    references = [
        '[1] IBM Security, \u201cCost of a Data Breach Report 2023,\u201d IBM Corporation, Armonk, NY, USA, 2023.',
        '[2] Ponemon Institute, \u201cThe Economics of Security Operations Centers: What is the True Cost for Effective Results?,\u201d Ponemon Institute LLC, 2020.',
        '[3] A. Chuvakin, K. Schmidt, and C. Phillips, Logging and Log Management: The Authoritative Guide to Understanding the Concepts Surrounding Logging and Log Management, Syngress, 2012.',
        '[4] R. Zuech, T. M. Khoshgoftaar, and R. Wald, \u201cIntrusion detection and big heterogeneous data: a survey,\u201d Journal of Big Data, vol. 2, no. 1, pp. 1\u201341, 2015.',
        '[5] M. Roesch, \u201cSnort \u2014 Lightweight intrusion detection for networks,\u201d in Proc. 13th LISA, USENIX, 1999, pp. 229\u2013238.',
        '[6] F. T. Liu, K. M. Ting, and Z.-H. Zhou, \u201cIsolation forest,\u201d in Proc. IEEE ICDM, 2008, pp. 413\u2013422.',
        '[7] L. Breiman, \u201cRandom forests,\u201d Machine Learning, vol. 45, no. 1, pp. 5\u201332, 2001.',
        '[8] I. Sharafaldin, A. H. Lashkari, and A. A. Ghorbani, \u201cToward generating a new intrusion detection dataset and intrusion traffic characterization,\u201d in Proc. ICISSP, 2018, pp. 108\u2013116.',
        '[9] N. Shone, T. N. Ngoc, V. D. Phai, and Q. Shi, \u201cA deep learning approach to network intrusion detection,\u201d IEEE Trans. Emerg. Topics Comput. Intell., vol. 2, no. 1, pp. 41\u201350, 2018.',
        '[10] M. Ferrag, O. Friha, D. Hamouda, L. Maglaras, and H. Janicke, \u201cEdge-IIoTset: A new comprehensive realistic cyber security dataset,\u201d IEEE Access, vol. 10, pp. 40281\u201340306, 2022.',
        '[11] M. Ahmed, A. N. Mahmood, and J. Hu, \u201cA survey of network anomaly detection techniques,\u201d J. Netw. Comput. Appl., vol. 60, pp. 19\u201331, 2016.',
        '[12] G. Apruzzese et al., \u201cOn the effectiveness of machine and deep learning for cyber security,\u201d in Proc. CyCon, 2018, pp. 371\u2013390.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        run.font.name = FONT_NAME
        run.font.size = REF_SIZE

    # ── Save ──
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "docs",
        "SOC_Copilot_IJIRT_Paper_v2.docx",
    )
    doc.save(output_path)
    print(f"Research paper saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_paper()
